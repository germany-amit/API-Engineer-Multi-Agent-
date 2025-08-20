import streamlit as st
import json, re
import PyPDF2
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from docx import Document
import nltk
from nltk.tokenize import word_tokenize
import os
from io import BytesIO

# --- Agentic AI Imports (Fail-Proof Cascade) ---
# Try to import LangChain
try:
    import langchain
    from langchain.prompts import PromptTemplate
    from langchain_huggingface import HuggingFaceHub
    from langchain.llms import HuggingFaceHub as LCHuggingFaceHub
    # Set framework to LangChain
    AGENT_FRAMEWORK = "LangChain"
except ImportError:
    # Try to import LlamaIndex
    try:
        import llama_index
        from llama_index.llms.huggingface import HuggingFaceLLM
        # Set framework to LlamaIndex
        AGENT_FRAMEWORK = "LlamaIndex"
    except ImportError:
        # Try to import AutoGen
        try:
            import autogen
            # Set framework to AutoGen
            AGENT_FRAMEWORK = "AutoGen"
        except ImportError:
            # Fallback to pure NLP/NLTK
            AGENT_FRAMEWORK = "NLTK"

# --- LLM and Model Cascade ---
LLM_MODEL = "placeholder"
llm_client = None

def get_llm():
    """Initializes LLM from Hugging Face Hub with a fail-proof cascade."""
    global LLM_MODEL, llm_client
    if os.environ.get("HUGGINGFACEHUB_API_TOKEN") is None:
        st.error("🚨 Hugging Face API token not set. Please add it to your Streamlit secrets.")
        st.stop()
    
    # Cascade list for free LLMs
    models_to_try = [
        "meta-llama/Llama-3-8B-Instruct",
        "mistralai/Mistral-7B-Instruct-v0.2",
        "google/gemma-7b-it",
    ]
    
    for model in models_to_try:
        try:
            if AGENT_FRAMEWORK == "LangChain":
                # Use langchain-huggingface for better integration
                llm = HuggingFaceHub(
                    repo_id=model, 
                    model_kwargs={"temperature": 0.1, "max_length": 1024}
                )
            elif AGENT_FRAMEWORK == "LlamaIndex":
                # Use LlamaIndex-specific LLM
                llm = HuggingFaceLLM(
                    model_name=model, 
                    max_new_tokens=1024,
                    model_kwargs={"temperature": 0.1}
                )
            else: # Fallback for AutoGen/NLTK
                 llm = LCHuggingFaceHub(
                    repo_id=model, 
                    model_kwargs={"temperature": 0.1, "max_length": 1024}
                )

            st.session_state.llm = llm
            LLM_MODEL = model
            st.success(f"✅ Using {AGENT_FRAMEWORK} with {LLM_MODEL}.")
            return llm
        except Exception:
            st.warning(f"❌ Failed to load {model}. Trying next model...")
            continue
    
    # Fallback to NLTK/Regex if no LLM works
    st.error("⚠️ Failed to load any LLM. Falling back to NLTK/Regex.")
    LLM_MODEL = "NLTK"
    st.session_state.llm = None
    return None

# --- Main App ---
st.title("📑 Agentic API Advisor from RFP")

# Call the LLM initializer once
if 'llm' not in st.session_state:
    get_llm()

# --- File Upload ---
uploaded_file = st.file_uploader("Upload RFP (txt/pdf)", type=["txt", "pdf"])

rfp_text = ""
if uploaded_file:
    if uploaded_file.type == "application/pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page in pdf_reader.pages:
            rfp_text += page.extract_text() or ""
    else: # TXT
        rfp_text = uploaded_file.read().decode("utf-8", errors="ignore")

    st.text_area("RFP Content", rfp_text, height=200)

# --- Generate Button ---
if st.button("🚀 Generate Agentic API Spec"):
    if not rfp_text:
        st.error("Please upload an RFP file first.")
    else:
        with st.spinner("🧠 Agent is thinking..."):
            if st.session_state.llm:
                # Use LangChain/LLM to generate the spec
                try:
                    prompt_template = PromptTemplate.from_template(
                        """You are a world-class API Engineer. Your task is to analyze the following RFP and generate a detailed API specification in JSON format.
                        
                        The JSON should include:
                        - "apiType": REST, GraphQL, or gRPC
                        - "endpoints": An array of API endpoints with brief descriptions.
                        - "security": The security mechanism required (e.g., OAuth2, API Keys).
                        - "dataModels": A JSON object defining the key data schemas.
                        
                        RFP Text:
                        {rfp_content}
                        
                        Your JSON response:
                        """
                    )
                    
                    if AGENT_FRAMEWORK == "LangChain":
                        chain = prompt_template | st.session_state.llm
                        result = chain.invoke({"rfp_content": rfp_text})
                    elif AGENT_FRAMEWORK == "LlamaIndex":
                         response = st.session_state.llm.complete(
                            prompt_template.format(rfp_content=rfp_text)
                        )
                         result = response.text
                    else: # AutoGen/NLTK fallback for logic
                         result = st.session_state.llm(
                            prompt_template.format(rfp_content=rfp_text)
                        )

                    # Extract JSON from the raw text response
                    json_match = re.search(r'```json\n(.*?)```', result, re.DOTALL)
                    if json_match:
                        spec_json = json_match.group(1).strip()
                    else:
                        spec_json = result.strip()
                    
                    try:
                        spec = json.loads(spec_json)
                        st.subheader("4️⃣ Draft API Spec")
                        st.json(spec)
                        st.session_state.spec = spec
                    except json.JSONDecodeError:
                        st.error("Failed to parse JSON from AI response. Displaying raw text instead.")
                        st.text(spec_json)
                        st.session_state.spec = None
                except Exception as e:
                    st.error(f"An error occurred during AI processing: {e}")
                    st.session_state.spec = None
            else: # NLTK/Regex fallback logic
                st.subheader("4️⃣ Draft API Spec (NLTK Fallback)")
                
                # Use NLTK to get a more nuanced spec
                try:
                    nltk.download('punkt', quiet=True)
                    tokens = word_tokenize(rfp_text.lower())
                except Exception:
                    tokens = rfp_text.lower().split()
                    
                spec = {
                    "apiType": "REST",
                    "endpoints": ["/sample"],
                    "security": "OAuth2 (placeholder)",
                    "dataModel": {}
                }

                if "graphql" in tokens: spec["apiType"] = "GraphQL"
                if any(x in tokens for x in ["real-time", "real time", "grpc"]): spec["apiType"] = "gRPC"
                if any(x in tokens for x in ["login", "token", "oauth"]): spec["security"] = "token based"
                if "user" in tokens: spec["dataModel"]["user"] = {"id": "integer", "name": "string"}
                
                st.json(spec)
                st.session_state.spec = spec

# --- Export ---
if 'spec' in st.session_state and st.session_state.spec:
    spec_content = json.dumps(st.session_state.spec, indent=2)

    col1, col2, col3 = st.columns(3)
    with col1:
        st.download_button("💾 Export as JSON", spec_content, file_name="api_spec.json", mime="application/json")
    with col2:
        output = BytesIO()
        pdf = SimpleDocTemplate(output)
        styles = getSampleStyleSheet()
        pdf.build([Paragraph(spec_content, styles['Normal'])])
        st.download_button("📄 Export as PDF", output.getvalue(), file_name="api_spec.pdf", mime="application/pdf")
    with col3:
        doc = Document()
        doc.add_heading("API Spec", 0)
        doc.add_paragraph(spec_content)
        doc_io = BytesIO()
        doc.save(doc_io)
        st.download_button("📝 Export as Word", doc_io.getvalue(), file_name="api_spec.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
